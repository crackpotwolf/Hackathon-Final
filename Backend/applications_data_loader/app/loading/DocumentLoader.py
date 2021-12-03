from typing import List
from docx import Document, table


class DataLoader:
  """Класс считывания данных из docs
  """

  def __init__(self):
    self.data_from_form = {}
    self.data_from_project_passport = {}

  def __load__(self, files):
    
    def get_data_from_vertical_table(table: table):
      """Метод считывания данных из  вертикальной таблицы

      Args:
          table (table): [текущая таблицы]

      Returns:
          [List]: [лист словарей, содержащие данные по ячейкам в формате название столбца: значение ячейки]
      """
      # метод парсинга таблицы
      data = []
      keys = None
      # проход по всем строкам таблицы
      for i, row in enumerate(table.rows):
        # забрать текст из ячейки
        text = [cell.text for cell in row.cells]
        # # если строка первая (0), то  задать заголовки столбцов
        # if i == 0:
        #     keys = tuple(text)
        #     continue
        # # запаковывание данных ячейки
        # row_data = dict(zip(keys, text))
        data.append(text)
      return data
    
    def get_data_from_horizontal_table(table: table):
      # метод парсинга таблицы
      data = []
      # проход по всем строкам таблицы
      for i, row in enumerate(table.rows):
        data.append({row.cells[0].text: row.cells[1].text})
      return data
    
    def get_data_from_text(document: Document):
      """Метод считываня текста из документа

      Args:
          document (Document): [открытый документ]

      Returns:
          [Dict]: [словарь с текстом из документа]
      """
      # получить все абзацы из документа
      paragraphs = document.paragraphs
      full_text = []
      # получить текст из каждого найденного параграфа
      for current_paragraph in paragraphs:
        full_text.append(current_paragraph.text)
      return full_text
    
    def check_document_for_form_tamplate(document_text: List):
      for paragraph in document_text:
        if "форма заявки" in paragraph.lower():
          return True
      return False
    
    def load_app_form(document_text: List):
      self.data_from_form = {
        "teamName": document_text[1].replace("1. Наименование команды/организации: ", ""),
        "stage": document_text[2].replace("2. Стадия готовности продукта: ", ""),
        "shortDescription": document_text[3].replace("3. Краткое описание продукта: ", ""),
        "cases": document_text[4].replace("4. Кейсы использования продукта: ", ""),
        "benefit": document_text[5].replace("5. Польза продукта: ", ""),
        "transportOrganization": document_text[6].replace("6. Организация Московского транспорта, интересная в первую очередь: ", ""),
        "pilotVision": document_text[7].replace("7. Запрос к акселератору и видение пилотного проекта: ", ""),
        "sertification": document_text[8].replace("8. Требуется ли сертификация продукта: ", ""),
        "contactFio": document_text[9].replace("9. ФИО контактного лица по заявке: ", ""),
        "contactPosition": document_text[10].replace("10. Должность контактного лица: ", ""),
        "phone": document_text[11].replace("11. Контактный телефон: ", ""),
        "email": document_text[12].replace("12. Контактная почта: ", ""),
        "legalName": document_text[13].replace("13. Наименование юридического лица: ", ""),
        "inn": document_text[14].replace("14. ИНН юридического лица: ", ""),
        "peopleCount": document_text[15].replace("15. Сколько человек в организации: ", ""),
        "site": document_text[16].replace("16. Сайт: ", ""),
        "acceleratorInfo": document_text[17].replace("17. Откуда узнали про акселератор: ", ""),
        "presentation": document_text[18].replace("18. Ссылка на презентацию: ", ""),
      }
    
    def load_product_passport(current_document: Document):
      #
      document_tables = current_document.tables
      #
      self.data_from_project_passport["passport"] = get_data_from_horizontal_table(table=document_tables[0])
      self.data_from_project_passport["context"] = document_tables[1].rows[0].cells[0].text
      self.data_from_project_passport["effects"] = get_data_from_vertical_table(table=document_tables[2])
      self.data_from_project_passport["stages"] = get_data_from_horizontal_table(table=document_tables[3])
      self.data_from_project_passport["teams"] = get_data_from_vertical_table(table=document_tables[4])
      self.data_from_project_passport["budget"] = get_data_from_vertical_table(table=document_tables[5])
      self.data_from_project_passport["statuses"] = get_data_from_vertical_table(table=document_tables[6])
      self.data_from_project_passport["activities"] = get_data_from_vertical_table(table=document_tables[7])
      self.data_from_project_passport["meetings"] = get_data_from_vertical_table(table=document_tables[8])
      self.data_from_project_passport["materials"] = get_data_from_vertical_table(table=document_tables[9])
    
    for file in files:
      #
      document = Document(file)
      
      #
      document_text = get_data_from_text(document=document)
      #
      if check_document_for_form_tamplate(document_text=document_text):
        load_app_form(document_text=document_text)
      else:
        load_product_passport(current_document=document)
      
