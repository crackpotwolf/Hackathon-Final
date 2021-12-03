from config import PATH_TO_RESULT_DATA
from docx import Document
from generating.ProductPassportGenerator import ProductPassportGenerator

import os


class ProductPassportDocxFormer:
  """Класс генерации файла Карточка проекта в формате *.docx
  """

  def __init__(self):
    self.created_status = False

  def __form_document__(self, product_passport_generator_num: int,  product_passport_generator: ProductPassportGenerator):
    document = Document()

    #
    title = document.add_paragraph('Паспорт проекта'.upper(), 'Title').bold = True

    #
    subtitle_1 = document.add_paragraph('Идентификация'.upper()).bold = True
    # добавляем таблицу
    table_1 = document.add_table(rows=8, cols=2)
    # применяем стиль для таблицы
    table_1.style = 'Table Grid'
    # заполняем таблицу данными
    table_1.cell(0, 0).text = "Наименование проекта"
    table_1.cell(0, 1).text = product_passport_generator.name
    table_1.cell(1, 0).text = "Организация транспортного комплекса Москвы"
    table_1.cell(1, 1).text = product_passport_generator.transportComplexOrganization
    table_1.cell(2, 0).text = "Участник программы пилотирования «Транспортные инновации Москвы»"
    table_1.cell(2, 1).text = product_passport_generator.pilotMember
    table_1.cell(3, 0).text = "Руководитель проекта"
    table_1.cell(3, 1).text = product_passport_generator.leader
    table_1.cell(4, 0).text = "Координатор от участника программы пилотирования «Транспортные инновации Москвы»"
    table_1.cell(4, 1).text = product_passport_generator.pilotCoordinator
    table_1.cell(5, 0).text = "Координатор от организации транспортного комплекса Москвы"
    table_1.cell(5, 1).text = product_passport_generator.transportComplexCoordinator
    table_1.cell(6, 0).text = "Краткое описание продукта"
    table_1.cell(6, 1).text = product_passport_generator.shortDescription
    table_1.cell(7, 0).text = "Сроки реализации проекта"
    table_1.cell(7, 1).text = product_passport_generator.timing

    #
    subtitle_2 = document.add_paragraph('Контекст и потребности'.upper()).bold = True
    # добавляем таблицу
    table_2 = document.add_table(rows=1, cols=1)
    # применяем стиль для таблицы
    table_2.style = 'Table Grid'
    # заполняем таблицу данными
    table_2.cell(0, 0).text = ""
    
    #
    subtitle_3 = document.add_paragraph('Ожидаемые эффекты проекта'.upper()).bold = True
    # добавляем таблицу
    table_3 = document.add_table(rows=len(product_passport_generator.effects), cols=3)
    # применяем стиль для таблицы
    table_3.style = 'Table Grid'
    # заполняем таблицу данными
    for num_row in range(len(product_passport_generator.effects)):
      table_3.cell(num_row, 0).text  = str(num_row+1)
      table_3.cell(num_row, 1).text  = product_passport_generator.effects[num_row]["name"]
      table_3.cell(num_row, 2).text  = product_passport_generator.effects[num_row]["value"]
    
    #
    subtitle_4 = document.add_paragraph('Этапы проекта'.upper()).bold = True
    # добавляем таблицу
    table_4 = document.add_table(rows=len(product_passport_generator.stages), cols=2)
    # применяем стиль для таблицы
    table_4.style = 'Table Grid'
    # заполняем таблицу данными
    for num_row in range(len(product_passport_generator.stages)):
      table_4.cell(num_row, 0).text  = product_passport_generator.stages[num_row]["name"]
      table_4.cell(num_row, 1).text  = product_passport_generator.stages[num_row]["date"]
    
    #
    subtitle_5 = document.add_paragraph('Команда проекта'.upper()).bold = True
    # добавляем таблицу
    table_5 = document.add_table(rows=len(product_passport_generator.teams), cols=4)
    # применяем стиль для таблицы
    table_5.style = 'Table Grid'
    # заполняем таблицу данными
    for num_row in range(len(product_passport_generator.teams)):
      table_5.cell(num_row, 0).text = str(num_row+1)
      table_5.cell(num_row, 1).text  = product_passport_generator.teams[num_row]["fio"]
      table_5.cell(num_row, 2).text  = product_passport_generator.teams[num_row]["position"]
      table_5.cell(num_row, 3).text  = product_passport_generator.teams[num_row]["contacts"]
    
    #
    subtitle_6 = document.add_paragraph('Бюджет проекта'.upper()).bold = True
    # добавляем таблицу
    table_6 = document.add_table(rows=len(product_passport_generator.budget), cols=3)
    # применяем стиль для таблицы
    table_6.style = 'Table Grid'
    # заполняем таблицу данными
    for num_row in range(len(product_passport_generator.budget)):
      table_6.cell(num_row, 0).text  = str(num_row+1)
      table_6.cell(num_row, 1).text  = product_passport_generator.budget[num_row]["name"]
      table_6.cell(num_row, 2).text  = product_passport_generator.budget[num_row]["value"]
    
    #
    subtitle_7 = document.add_paragraph('Статус проекта'.upper(), "Title").bold = True
    # добавляем таблицу
    table_7 = document.add_table(rows=len(product_passport_generator.statuses), cols=3)
    # применяем стиль для таблицы
    table_7.style = 'Table Grid'
    # заполняем таблицу данными
    for num_row in range(len(product_passport_generator.statuses)):
      table_7.cell(num_row, 0).text  = product_passport_generator.statuses[num_row]["name"]
      table_7.cell(num_row, 1).text  = product_passport_generator.statuses[num_row]["expectations"]
      table_7.cell(num_row, 2).text  = product_passport_generator.statuses[num_row]["date"]
    
    #
    subtitle_8 = document.add_paragraph('Мероприятия по проекту'.upper(), "Title").bold = True
    # добавляем таблицу
    table_8 = document.add_table(
        rows=len(product_passport_generator.activities), cols=4)
    # применяем стиль для таблицы
    table_8.style = 'Table Grid'
    # заполняем таблицу данными
    for num_row in range(len(product_passport_generator.activities)):
      table_8.cell(num_row, 0).text = str(num_row+1)
      table_8.cell(num_row, 1).text  = product_passport_generator.activities[num_row]["name"]
      table_8.cell(num_row, 1).text  = product_passport_generator.activities[num_row]["expectations"]
      table_8.cell(num_row, 2).text  = product_passport_generator.activities[num_row]["date"]
    
     #
    subtitle_9 = document.add_paragraph('Собрания по проекту'.upper(), "Title").bold = True
    # добавляем таблицу
    table_9  = document.add_table(rows=len(product_passport_generator.meetings), cols=2)
    # применяем стиль для таблицы
    table_9.style = 'Table Grid'
    # заполняем таблицу данными
    for num_row in range(len(product_passport_generator.meetings)):
      table_9.cell(num_row, 0).text  = str(num_row+1)
      table_9.cell(num_row, 1).text  = product_passport_generator.meetings[num_row]["comment"]
      table_9.cell(num_row, 1).text  = product_passport_generator.meetings[num_row]["date"]
    
    #
    subtitle_10 = document.add_paragraph('Материалы проекта'.upper()).bold = True
    # добавляем таблицу
    table_10 = document.add_table(rows=len(product_passport_generator.materials), cols=2)
    # применяем стиль для таблицы
    table_10.style = 'Table Grid'
    # заполняем таблицу данными
    for num_row in range(len(product_passport_generator.materials)):
      table_10.cell(num_row, 0).text  = product_passport_generator.materials[num_row]["name"]
      table_10.cell(num_row, 1).text  = product_passport_generator.materials[num_row]["link"]
      
    #
    filename = os.path.join(PATH_TO_RESULT_DATA, 'Паспорт проекта %s.docx' % (str(product_passport_generator_num+1)))
    document.save(filename)

    self.created_status = True
