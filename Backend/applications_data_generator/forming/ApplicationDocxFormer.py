from config import PATH_TO_RESULT_DATA
from docx import Document
from generating.ApplicationGenerator import ApplicationGenerator

import os


class ApplicationDocxFormer:
  """Класс генерации файла Формы заявки в формате *.docx
  """
  def __init__(self):
    self.created_status = False
    
  def __form_document__(self, application_generator_num: int, application_generator: ApplicationGenerator):
    document = Document()
    
    #
    title = document.add_paragraph('Форма заявки для стартапа: ', 'Title').bold = True
    
    # добавляем параграфы
    prgrph_1 = document.add_paragraph("1. Наименование команды/организации: %s" % (application_generator.fullname_participant))
    prgrph_2 = document.add_paragraph("2. Стадия готовности продукта: %s" % (application_generator.stage_product))
    prgrph_3 = document.add_paragraph("3. Краткое описание продукта: %s" % (application_generator.short_description))
    prgrph_4 = document.add_paragraph("4. Кейсы использования продукта: %s" % (application_generator.cases_product_using))
    prgrph_5 = document.add_paragraph("5. Польза продукта: %s" % (application_generator.product_use))
    prgrph_6 = document.add_paragraph("6. Организация Московского транспорта, интересная в первую очередь: %s" % (application_generator.interest_transport_organization))
    prgrph_7 = document.add_paragraph("7. Запрос к акселератору и видение пилотного проекта: %s" % (application_generator.request_to_accelerate))
    prgrph_8 = document.add_paragraph("8. Требуется ли сертификация продукта: %s" % (application_generator.certification))
    prgrph_9 = document.add_paragraph("9. ФИО контактного лица по заявке: %s" % (application_generator.fullname_applicant))
    prgrph_0 = document.add_paragraph("10. Должность контактного лица: %s" % (application_generator.role_applicant))
    prgrph_11 = document.add_paragraph("11. Контактный телефон: %s" % (application_generator.phone_applicant))
    prgrph_12 = document.add_paragraph("12. Контактная почта: %s" % (application_generator.email_applicant))
    prgrph_13 = document.add_paragraph("13. Наименование юридического лица: %s" % (application_generator.name_company))
    prgrph_14 = document.add_paragraph("14. ИНН юридического лица: %s" % (application_generator.inn_company))
    prgrph_15 = document.add_paragraph("15. Сколько человек в организации: %s" % (application_generator.num_people_company))
    prgrph_16 = document.add_paragraph("16. Сайт: %s" % (application_generator.site_company))
    prgrph_17 = document.add_paragraph("17. Откуда узнали про акселератор: %s" % (application_generator.source_info_accelerator))
    prgrph_18 = document.add_paragraph("18. Ссылка на презентацию: %s" % (application_generator.link_presentation))
    #
    filename = os.path.join(PATH_TO_RESULT_DATA, 'Форма заявки для стартапа %s.docx' % (str(application_generator_num+1)))
    document.save(filename)

    self.created_status = True
